"""
Group Coordination and Management System for Spirit Tours
Handles group organization, guide assignment, vouchers, and comprehensive reporting
"""

import asyncio
import json
from datetime import datetime, timedelta, date
from typing import Dict, List, Optional, Any, Union
from enum import Enum
from dataclasses import dataclass, field, asdict
import uuid
from collections import defaultdict
import pandas as pd
import numpy as np
from sqlalchemy import select, and_, or_, func, text, update, delete
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload, joinedload
import aiohttp
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
import smtplib
import ssl
from jinja2 import Template
import pytz
from redis import asyncio as aioredis
import hashlib
import pickle
from functools import wraps
import logging
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import io
import base64
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import openpyxl
from openpyxl.styles import Font, Fill, PatternFill, Alignment, Border, Side
import qrcode
from barcode import Code128
from barcode.writer import ImageWriter

logger = logging.getLogger(__name__)

# Enums for system states
class GroupStatus(Enum):
    DRAFT = "draft"
    PENDING_ASSIGNMENT = "pending_assignment"
    PARTIALLY_ASSIGNED = "partially_assigned"
    FULLY_ASSIGNED = "fully_assigned"
    CONFIRMED = "confirmed"
    IN_PROGRESS = "in_progress"
    COMPLETED = "completed"
    CANCELLED = "cancelled"

class VoucherType(Enum):
    HOTEL = "hotel"
    RESTAURANT = "restaurant"
    ENTRANCE_TICKET = "entrance_ticket"
    TRANSPORT = "transport"
    ACTIVITY = "activity"
    OTHER = "other"

class VoucherStatus(Enum):
    PENDING = "pending"
    CONFIRMED = "confirmed"
    USED = "used"
    CANCELLED = "cancelled"
    EXPIRED = "expired"

class AssignmentStatus(Enum):
    NOT_ASSIGNED = "not_assigned"
    PENDING_CONFIRMATION = "pending_confirmation"
    CONFIRMED = "confirmed"
    BACKUP_ASSIGNED = "backup_assigned"

class ReminderFrequency(Enum):
    BIWEEKLY = "biweekly"
    EVERY_3_DAYS = "every_3_days"
    DAILY = "daily"
    URGENT_DAILY = "urgent_daily"

class ReportFormat(Enum):
    PDF = "pdf"
    EXCEL = "excel"
    WORD = "word"
    HTML = "html"
    JSON = "json"

# Data Models
@dataclass
class ContactInfo:
    """Contact information for guides, drivers, coordinators"""
    primary_phone: str
    secondary_phone: Optional[str] = None
    email: Optional[str] = None
    whatsapp: Optional[str] = None
    emergency_contact: Optional[str] = None
    emergency_phone: Optional[str] = None

@dataclass
class GuideAssignment:
    """Guide/driver/coordinator assignment details"""
    assignment_id: str = field(default_factory=lambda: str(uuid.uuid4()))
    group_id: str = ""
    role: str = ""  # guide, driver, coordinator
    person_id: str = ""
    person_name: str = ""
    contact_info: Optional[ContactInfo] = None
    status: AssignmentStatus = AssignmentStatus.NOT_ASSIGNED
    assigned_date: Optional[datetime] = None
    confirmed_date: Optional[datetime] = None
    notes: Optional[str] = None
    backup_person_id: Optional[str] = None
    backup_contact: Optional[ContactInfo] = None

@dataclass
class Voucher:
    """Voucher information for hotels, restaurants, entrances"""
    voucher_id: str = field(default_factory=lambda: str(uuid.uuid4()))
    voucher_number: str = ""
    group_id: str = ""
    type: VoucherType = VoucherType.OTHER
    status: VoucherStatus = VoucherStatus.PENDING
    provider_name: str = ""
    provider_contact: Optional[str] = None
    service_date: Optional[date] = None
    service_time: Optional[str] = None
    quantity: int = 1
    unit_price: float = 0.0
    total_amount: float = 0.0
    currency: str = "USD"
    confirmation_code: Optional[str] = None
    special_requirements: Optional[str] = None
    qr_code: Optional[str] = None
    barcode: Optional[str] = None
    created_date: datetime = field(default_factory=datetime.utcnow)
    expiry_date: Optional[datetime] = None
    used_date: Optional[datetime] = None
    metadata: Dict[str, Any] = field(default_factory=dict)

@dataclass
class HotelVoucher(Voucher):
    """Specific voucher for hotel bookings"""
    check_in_date: Optional[date] = None
    check_out_date: Optional[date] = None
    nights: int = 0
    room_type: Optional[str] = None
    meal_plan: Optional[str] = None
    rooms_count: int = 1
    guests_per_room: int = 2
    total_guests: int = 2
    rooming_list: List[Dict[str, str]] = field(default_factory=list)

@dataclass
class RestaurantVoucher(Voucher):
    """Specific voucher for restaurant bookings"""
    meal_type: Optional[str] = None  # breakfast, lunch, dinner
    number_of_guests: int = 1
    menu_type: Optional[str] = None
    dietary_requirements: List[str] = field(default_factory=list)
    table_preference: Optional[str] = None

@dataclass
class EntranceVoucher(Voucher):
    """Specific voucher for entrance tickets"""
    attraction_name: str = ""
    ticket_type: Optional[str] = None  # adult, child, senior
    valid_from: Optional[datetime] = None
    valid_until: Optional[datetime] = None
    number_of_tickets: int = 1
    includes_guide: bool = False
    includes_transport: bool = False

@dataclass
class FlightInfo:
    """Flight information for group travel"""
    flight_number: str
    airline: str
    departure_airport: str
    arrival_airport: str
    departure_time: datetime
    arrival_time: datetime
    terminal: Optional[str] = None
    gate: Optional[str] = None
    passengers: List[Dict[str, Any]] = field(default_factory=list)
    booking_reference: Optional[str] = None
    seat_assignments: Dict[str, str] = field(default_factory=dict)

@dataclass
class GroupCoordination:
    """Main group coordination entity"""
    group_id: str = field(default_factory=lambda: str(uuid.uuid4()))
    group_number: str = ""
    group_name: str = ""
    status: GroupStatus = GroupStatus.DRAFT
    travel_date_from: Optional[date] = None
    travel_date_to: Optional[date] = None
    total_participants: int = 0
    group_type: str = "mixed"  # individual, group, reserved
    
    # Assignments
    guide_assignment: Optional[GuideAssignment] = None
    driver_assignment: Optional[GuideAssignment] = None
    coordinator_assignment: Optional[GuideAssignment] = None
    
    # Contact Information
    emergency_contacts: List[ContactInfo] = field(default_factory=list)
    
    # Travel Information
    flights: List[FlightInfo] = field(default_factory=list)
    
    # Vouchers
    hotel_vouchers: List[HotelVoucher] = field(default_factory=list)
    restaurant_vouchers: List[RestaurantVoucher] = field(default_factory=list)
    entrance_vouchers: List[EntranceVoucher] = field(default_factory=list)
    other_vouchers: List[Voucher] = field(default_factory=list)
    
    # Participants
    participants: List[Dict[str, Any]] = field(default_factory=list)
    rooming_list: List[Dict[str, Any]] = field(default_factory=list)
    
    # Metadata
    created_date: datetime = field(default_factory=datetime.utcnow)
    last_updated: datetime = field(default_factory=datetime.utcnow)
    created_by: Optional[str] = None
    notes: Optional[str] = None
    internal_notes: Optional[str] = None
    
    # Reminder Settings
    reminder_enabled: bool = True
    last_reminder_sent: Optional[datetime] = None
    reminder_frequency: Optional[ReminderFrequency] = None
    custom_reminder_emails: List[str] = field(default_factory=list)

@dataclass
class ReportConfiguration:
    """Configuration for customizable reports"""
    report_id: str = field(default_factory=lambda: str(uuid.uuid4()))
    report_name: str = ""
    report_type: str = "group_complete"  # group_complete, rooming_list, flight_manifest, voucher_summary
    format: ReportFormat = ReportFormat.PDF
    
    # Sections to include
    include_group_info: bool = True
    include_participants: bool = True
    include_flights: bool = True
    include_hotels: bool = True
    include_rooming_list: bool = True
    include_restaurants: bool = True
    include_entrances: bool = True
    include_assignments: bool = True
    include_vouchers: bool = True
    include_emergency_contacts: bool = True
    include_notes: bool = False
    include_costs: bool = False
    include_qr_codes: bool = True
    include_barcodes: bool = False
    
    # Grouping and Sorting
    group_by: Optional[str] = None  # flight, hotel, date, service_type
    sort_by: Optional[str] = None  # name, date, flight, room
    
    # Filtering
    date_from: Optional[date] = None
    date_to: Optional[date] = None
    status_filter: Optional[List[str]] = None
    service_filter: Optional[List[str]] = None
    
    # Formatting
    paper_size: str = "A4"  # A4, Letter, Legal
    orientation: str = "portrait"  # portrait, landscape
    font_size: int = 10
    include_logo: bool = True
    include_header: bool = True
    include_footer: bool = True
    include_page_numbers: bool = True
    
    # Distribution
    email_to: List[str] = field(default_factory=list)
    auto_send: bool = False
    schedule_time: Optional[datetime] = None

class GroupCoordinationSystem:
    """Main system for managing group coordination, assignments, and vouchers"""
    
    def __init__(self, db_session: AsyncSession, redis_client: aioredis.Redis = None):
        self.db = db_session
        self.redis = redis_client
        self.cache_ttl = 3600  # 1 hour cache
        self.reminder_queue = asyncio.Queue()
        self.report_queue = asyncio.Queue()
        self.notification_service = NotificationService()
        self.voucher_generator = VoucherGenerator()
        self.report_generator = ReportGenerator()
        
    async def create_group_coordination(self, group_data: Dict[str, Any]) -> GroupCoordination:
        """Create a new group coordination entity"""
        try:
            # Generate unique group number
            group_number = await self._generate_group_number()
            
            # Create group coordination
            group = GroupCoordination(
                group_number=group_number,
                group_name=group_data.get('group_name', ''),
                status=GroupStatus.DRAFT,
                travel_date_from=group_data.get('travel_date_from'),
                travel_date_to=group_data.get('travel_date_to'),
                total_participants=group_data.get('total_participants', 0),
                group_type=group_data.get('group_type', 'mixed'),
                created_by=group_data.get('created_by'),
                notes=group_data.get('notes'),
                reminder_enabled=group_data.get('reminder_enabled', True),
                custom_reminder_emails=group_data.get('reminder_emails', [])
            )
            
            # Add participants if provided
            if 'participants' in group_data:
                group.participants = group_data['participants']
            
            # Add flights if provided
            if 'flights' in group_data:
                for flight_data in group_data['flights']:
                    flight = FlightInfo(**flight_data)
                    group.flights.append(flight)
            
            # Save to database
            await self._save_group_coordination(group)
            
            # Start reminder monitoring
            if group.reminder_enabled:
                asyncio.create_task(self._monitor_group_reminders(group.group_id))
            
            # Cache the group
            await self._cache_group(group)
            
            logger.info(f"Created group coordination: {group.group_number}")
            return group
            
        except Exception as e:
            logger.error(f"Error creating group coordination: {str(e)}")
            raise
    
    async def assign_guide(self, group_id: str, guide_data: Dict[str, Any]) -> GuideAssignment:
        """Assign a guide to a group"""
        try:
            group = await self.get_group_coordination(group_id)
            if not group:
                raise ValueError(f"Group {group_id} not found")
            
            # Create guide assignment
            assignment = GuideAssignment(
                group_id=group_id,
                role="guide",
                person_id=guide_data.get('person_id'),
                person_name=guide_data.get('person_name'),
                contact_info=ContactInfo(**guide_data.get('contact_info', {})),
                status=AssignmentStatus.PENDING_CONFIRMATION,
                assigned_date=datetime.utcnow(),
                notes=guide_data.get('notes')
            )
            
            # Add backup if provided
            if 'backup_person_id' in guide_data:
                assignment.backup_person_id = guide_data['backup_person_id']
                if 'backup_contact' in guide_data:
                    assignment.backup_contact = ContactInfo(**guide_data['backup_contact'])
            
            # Update group
            group.guide_assignment = assignment
            await self._update_group_status(group)
            await self._save_group_coordination(group)
            
            # Send confirmation request
            await self._send_assignment_confirmation(assignment)
            
            # Clear reminders if all assignments are complete
            await self._check_and_clear_reminders(group)
            
            logger.info(f"Assigned guide {assignment.person_name} to group {group.group_number}")
            return assignment
            
        except Exception as e:
            logger.error(f"Error assigning guide: {str(e)}")
            raise
    
    async def assign_driver(self, group_id: str, driver_data: Dict[str, Any]) -> GuideAssignment:
        """Assign a driver to a group"""
        try:
            group = await self.get_group_coordination(group_id)
            if not group:
                raise ValueError(f"Group {group_id} not found")
            
            # Create driver assignment
            assignment = GuideAssignment(
                group_id=group_id,
                role="driver",
                person_id=driver_data.get('person_id'),
                person_name=driver_data.get('person_name'),
                contact_info=ContactInfo(**driver_data.get('contact_info', {})),
                status=AssignmentStatus.PENDING_CONFIRMATION,
                assigned_date=datetime.utcnow(),
                notes=driver_data.get('notes')
            )
            
            # Update group
            group.driver_assignment = assignment
            await self._update_group_status(group)
            await self._save_group_coordination(group)
            
            logger.info(f"Assigned driver {assignment.person_name} to group {group.group_number}")
            return assignment
            
        except Exception as e:
            logger.error(f"Error assigning driver: {str(e)}")
            raise
    
    async def create_hotel_voucher(self, group_id: str, hotel_data: Dict[str, Any]) -> HotelVoucher:
        """Create a hotel voucher for a group"""
        try:
            group = await self.get_group_coordination(group_id)
            if not group:
                raise ValueError(f"Group {group_id} not found")
            
            # Generate voucher number
            voucher_number = await self.voucher_generator.generate_voucher_number(VoucherType.HOTEL)
            
            # Create hotel voucher
            voucher = HotelVoucher(
                voucher_number=voucher_number,
                group_id=group_id,
                type=VoucherType.HOTEL,
                status=VoucherStatus.PENDING,
                provider_name=hotel_data.get('hotel_name'),
                provider_contact=hotel_data.get('hotel_contact'),
                check_in_date=hotel_data.get('check_in_date'),
                check_out_date=hotel_data.get('check_out_date'),
                nights=hotel_data.get('nights', 1),
                room_type=hotel_data.get('room_type'),
                meal_plan=hotel_data.get('meal_plan'),
                rooms_count=hotel_data.get('rooms_count', 1),
                guests_per_room=hotel_data.get('guests_per_room', 2),
                total_guests=hotel_data.get('total_guests'),
                rooming_list=hotel_data.get('rooming_list', []),
                confirmation_code=hotel_data.get('confirmation_code'),
                special_requirements=hotel_data.get('special_requirements'),
                total_amount=hotel_data.get('total_amount', 0),
                currency=hotel_data.get('currency', 'USD')
            )
            
            # Generate QR code for voucher
            voucher.qr_code = await self.voucher_generator.generate_qr_code(voucher)
            
            # Add to group
            group.hotel_vouchers.append(voucher)
            await self._save_group_coordination(group)
            
            logger.info(f"Created hotel voucher {voucher_number} for group {group.group_number}")
            return voucher
            
        except Exception as e:
            logger.error(f"Error creating hotel voucher: {str(e)}")
            raise
    
    async def create_restaurant_voucher(self, group_id: str, restaurant_data: Dict[str, Any]) -> RestaurantVoucher:
        """Create a restaurant voucher for a group"""
        try:
            group = await self.get_group_coordination(group_id)
            if not group:
                raise ValueError(f"Group {group_id} not found")
            
            # Generate voucher number
            voucher_number = await self.voucher_generator.generate_voucher_number(VoucherType.RESTAURANT)
            
            # Create restaurant voucher
            voucher = RestaurantVoucher(
                voucher_number=voucher_number,
                group_id=group_id,
                type=VoucherType.RESTAURANT,
                status=VoucherStatus.PENDING,
                provider_name=restaurant_data.get('restaurant_name'),
                provider_contact=restaurant_data.get('restaurant_contact'),
                service_date=restaurant_data.get('service_date'),
                service_time=restaurant_data.get('service_time'),
                meal_type=restaurant_data.get('meal_type'),
                number_of_guests=restaurant_data.get('number_of_guests', 1),
                menu_type=restaurant_data.get('menu_type'),
                dietary_requirements=restaurant_data.get('dietary_requirements', []),
                table_preference=restaurant_data.get('table_preference'),
                confirmation_code=restaurant_data.get('confirmation_code'),
                special_requirements=restaurant_data.get('special_requirements'),
                total_amount=restaurant_data.get('total_amount', 0),
                currency=restaurant_data.get('currency', 'USD')
            )
            
            # Generate QR code for voucher
            voucher.qr_code = await self.voucher_generator.generate_qr_code(voucher)
            
            # Add to group
            group.restaurant_vouchers.append(voucher)
            await self._save_group_coordination(group)
            
            logger.info(f"Created restaurant voucher {voucher_number} for group {group.group_number}")
            return voucher
            
        except Exception as e:
            logger.error(f"Error creating restaurant voucher: {str(e)}")
            raise
    
    async def create_entrance_voucher(self, group_id: str, entrance_data: Dict[str, Any]) -> EntranceVoucher:
        """Create an entrance ticket voucher for a group"""
        try:
            group = await self.get_group_coordination(group_id)
            if not group:
                raise ValueError(f"Group {group_id} not found")
            
            # Generate voucher number
            voucher_number = await self.voucher_generator.generate_voucher_number(VoucherType.ENTRANCE_TICKET)
            
            # Create entrance voucher
            voucher = EntranceVoucher(
                voucher_number=voucher_number,
                group_id=group_id,
                type=VoucherType.ENTRANCE_TICKET,
                status=VoucherStatus.PENDING,
                provider_name=entrance_data.get('provider_name'),
                attraction_name=entrance_data.get('attraction_name'),
                service_date=entrance_data.get('visit_date'),
                service_time=entrance_data.get('visit_time'),
                ticket_type=entrance_data.get('ticket_type'),
                valid_from=entrance_data.get('valid_from'),
                valid_until=entrance_data.get('valid_until'),
                number_of_tickets=entrance_data.get('number_of_tickets', 1),
                includes_guide=entrance_data.get('includes_guide', False),
                includes_transport=entrance_data.get('includes_transport', False),
                confirmation_code=entrance_data.get('confirmation_code'),
                special_requirements=entrance_data.get('special_requirements'),
                total_amount=entrance_data.get('total_amount', 0),
                currency=entrance_data.get('currency', 'USD')
            )
            
            # Generate QR code and barcode for voucher
            voucher.qr_code = await self.voucher_generator.generate_qr_code(voucher)
            voucher.barcode = await self.voucher_generator.generate_barcode(voucher)
            
            # Add to group
            group.entrance_vouchers.append(voucher)
            await self._save_group_coordination(group)
            
            logger.info(f"Created entrance voucher {voucher_number} for group {group.group_number}")
            return voucher
            
        except Exception as e:
            logger.error(f"Error creating entrance voucher: {str(e)}")
            raise
    
    async def generate_group_report(self, group_id: str, config: ReportConfiguration) -> bytes:
        """Generate a customizable report for a group"""
        try:
            group = await self.get_group_coordination(group_id)
            if not group:
                raise ValueError(f"Group {group_id} not found")
            
            # Generate report based on configuration
            report_data = await self._prepare_report_data(group, config)
            
            # Generate report in requested format
            if config.format == ReportFormat.PDF:
                report = await self.report_generator.generate_pdf_report(report_data, config)
            elif config.format == ReportFormat.EXCEL:
                report = await self.report_generator.generate_excel_report(report_data, config)
            elif config.format == ReportFormat.WORD:
                report = await self.report_generator.generate_word_report(report_data, config)
            elif config.format == ReportFormat.HTML:
                report = await self.report_generator.generate_html_report(report_data, config)
            else:
                report = await self.report_generator.generate_json_report(report_data, config)
            
            # Auto-send if configured
            if config.auto_send and config.email_to:
                await self._send_report_email(report, config, group)
            
            logger.info(f"Generated {config.format.value} report for group {group.group_number}")
            return report
            
        except Exception as e:
            logger.error(f"Error generating group report: {str(e)}")
            raise
    
    async def get_rooming_list_report(self, group_id: str, hotel_id: Optional[str] = None) -> Dict[str, Any]:
        """Generate a rooming list report for a group"""
        try:
            group = await self.get_group_coordination(group_id)
            if not group:
                raise ValueError(f"Group {group_id} not found")
            
            rooming_data = {
                'group_number': group.group_number,
                'group_name': group.group_name,
                'hotels': []
            }
            
            # Filter hotels if specific hotel requested
            hotels = group.hotel_vouchers
            if hotel_id:
                hotels = [h for h in hotels if h.voucher_id == hotel_id]
            
            for hotel in hotels:
                hotel_info = {
                    'hotel_name': hotel.provider_name,
                    'check_in': hotel.check_in_date.isoformat() if hotel.check_in_date else None,
                    'check_out': hotel.check_out_date.isoformat() if hotel.check_out_date else None,
                    'nights': hotel.nights,
                    'total_rooms': hotel.rooms_count,
                    'total_guests': hotel.total_guests,
                    'rooming_list': hotel.rooming_list
                }
                rooming_data['hotels'].append(hotel_info)
            
            return rooming_data
            
        except Exception as e:
            logger.error(f"Error generating rooming list: {str(e)}")
            raise
    
    async def get_flight_manifest_report(self, group_id: str, flight_number: Optional[str] = None) -> Dict[str, Any]:
        """Generate a flight manifest report for a group"""
        try:
            group = await self.get_group_coordination(group_id)
            if not group:
                raise ValueError(f"Group {group_id} not found")
            
            manifest_data = {
                'group_number': group.group_number,
                'group_name': group.group_name,
                'flights': []
            }
            
            # Filter flights if specific flight requested
            flights = group.flights
            if flight_number:
                flights = [f for f in flights if f.flight_number == flight_number]
            
            for flight in flights:
                # Organize passengers by flight
                flight_info = {
                    'flight_number': flight.flight_number,
                    'airline': flight.airline,
                    'departure': {
                        'airport': flight.departure_airport,
                        'time': flight.departure_time.isoformat(),
                        'terminal': flight.terminal,
                        'gate': flight.gate
                    },
                    'arrival': {
                        'airport': flight.arrival_airport,
                        'time': flight.arrival_time.isoformat()
                    },
                    'booking_reference': flight.booking_reference,
                    'passengers': flight.passengers,
                    'seat_assignments': flight.seat_assignments,
                    'total_passengers': len(flight.passengers)
                }
                manifest_data['flights'].append(flight_info)
            
            return manifest_data
            
        except Exception as e:
            logger.error(f"Error generating flight manifest: {str(e)}")
            raise
    
    async def check_missing_assignments(self, group_id: str) -> Dict[str, Any]:
        """Check for missing assignments in a group"""
        try:
            group = await self.get_group_coordination(group_id)
            if not group:
                raise ValueError(f"Group {group_id} not found")
            
            missing = {
                'group_id': group_id,
                'group_number': group.group_number,
                'days_until_travel': None,
                'missing_assignments': [],
                'incomplete_vouchers': [],
                'missing_data': []
            }
            
            # Calculate days until travel
            if group.travel_date_from:
                days_until = (group.travel_date_from - date.today()).days
                missing['days_until_travel'] = days_until
            
            # Check guide assignment
            if not group.guide_assignment or group.guide_assignment.status != AssignmentStatus.CONFIRMED:
                missing['missing_assignments'].append({
                    'type': 'guide',
                    'status': group.guide_assignment.status.value if group.guide_assignment else 'not_assigned'
                })
            
            # Check driver assignment
            if not group.driver_assignment or group.driver_assignment.status != AssignmentStatus.CONFIRMED:
                missing['missing_assignments'].append({
                    'type': 'driver',
                    'status': group.driver_assignment.status.value if group.driver_assignment else 'not_assigned'
                })
            
            # Check coordinator assignment
            if not group.coordinator_assignment or group.coordinator_assignment.status != AssignmentStatus.CONFIRMED:
                missing['missing_assignments'].append({
                    'type': 'coordinator',
                    'status': group.coordinator_assignment.status.value if group.coordinator_assignment else 'not_assigned'
                })
            
            # Check voucher confirmations
            for voucher in group.hotel_vouchers:
                if voucher.status != VoucherStatus.CONFIRMED:
                    missing['incomplete_vouchers'].append({
                        'type': 'hotel',
                        'voucher_number': voucher.voucher_number,
                        'status': voucher.status.value
                    })
            
            for voucher in group.restaurant_vouchers:
                if voucher.status != VoucherStatus.CONFIRMED:
                    missing['incomplete_vouchers'].append({
                        'type': 'restaurant',
                        'voucher_number': voucher.voucher_number,
                        'status': voucher.status.value
                    })
            
            # Check for missing contact information
            if group.guide_assignment and not group.guide_assignment.contact_info:
                missing['missing_data'].append('Guide contact information')
            
            if group.driver_assignment and not group.driver_assignment.contact_info:
                missing['missing_data'].append('Driver contact information')
            
            # Check for missing rooming lists
            for hotel in group.hotel_vouchers:
                if not hotel.rooming_list:
                    missing['missing_data'].append(f'Rooming list for {hotel.provider_name}')
            
            return missing
            
        except Exception as e:
            logger.error(f"Error checking missing assignments: {str(e)}")
            raise
    
    async def _monitor_group_reminders(self, group_id: str):
        """Monitor and send reminders for missing group data"""
        try:
            while True:
                group = await self.get_group_coordination(group_id)
                if not group or not group.reminder_enabled:
                    break
                
                if not group.travel_date_from:
                    await asyncio.sleep(86400)  # Check daily if no travel date
                    continue
                
                days_until_travel = (group.travel_date_from - date.today()).days
                
                # Check for missing data
                missing_data = await self.check_missing_assignments(group_id)
                
                if missing_data['missing_assignments'] or missing_data['incomplete_vouchers'] or missing_data['missing_data']:
                    # Determine reminder frequency based on urgency
                    if days_until_travel <= 14:  # 2 weeks or less
                        frequency = ReminderFrequency.URGENT_DAILY
                        sleep_seconds = 86400  # Daily
                    elif days_until_travel <= 30:  # 1 month or less
                        frequency = ReminderFrequency.EVERY_3_DAYS
                        sleep_seconds = 259200  # Every 3 days
                    else:  # More than 1 month
                        frequency = ReminderFrequency.BIWEEKLY
                        sleep_seconds = 1209600  # Every 2 weeks
                    
                    # Check if reminder should be sent
                    should_send = False
                    if not group.last_reminder_sent:
                        should_send = True
                    else:
                        time_since_last = (datetime.utcnow() - group.last_reminder_sent).total_seconds()
                        if time_since_last >= sleep_seconds:
                            should_send = True
                    
                    if should_send:
                        # Send reminder
                        await self._send_missing_data_reminder(group, missing_data, frequency)
                        
                        # Update last reminder sent
                        group.last_reminder_sent = datetime.utcnow()
                        group.reminder_frequency = frequency
                        await self._save_group_coordination(group)
                
                # Sleep until next check
                await asyncio.sleep(3600)  # Check every hour
                
        except Exception as e:
            logger.error(f"Error in reminder monitoring for group {group_id}: {str(e)}")
    
    async def _send_missing_data_reminder(self, group: GroupCoordination, missing_data: Dict[str, Any], frequency: ReminderFrequency):
        """Send reminder email about missing data"""
        try:
            # Determine urgency level
            urgency = "STANDARD"
            subject_prefix = "Reminder"
            if frequency == ReminderFrequency.URGENT_DAILY:
                urgency = "URGENT"
                subject_prefix = "⚠️ URGENT"
            elif frequency == ReminderFrequency.EVERY_3_DAYS:
                urgency = "IMPORTANT"
                subject_prefix = "⚠️ IMPORTANT"
            
            # Prepare email content
            subject = f"{subject_prefix}: Missing Data for Group {group.group_number} - {group.group_name}"
            
            # Build email body
            body = f"""
            <html>
            <body style="font-family: Arial, sans-serif;">
                <h2 style="color: {'red' if urgency == 'URGENT' else 'orange'};">
                    {subject_prefix}: Action Required for Group {group.group_number}
                </h2>
                
                <p><strong>Group Name:</strong> {group.group_name}</p>
                <p><strong>Travel Date:</strong> {group.travel_date_from} to {group.travel_date_to}</p>
                <p><strong>Days Until Travel:</strong> {missing_data['days_until_travel']} days</p>
                
                <h3>Missing Assignments:</h3>
                <ul>
            """
            
            for assignment in missing_data['missing_assignments']:
                body += f"<li style='color: red;'>❌ {assignment['type'].upper()}: {assignment['status']}</li>"
            
            if not missing_data['missing_assignments']:
                body += "<li style='color: green;'>✅ All assignments complete</li>"
            
            body += """
                </ul>
                
                <h3>Incomplete Vouchers:</h3>
                <ul>
            """
            
            for voucher in missing_data['incomplete_vouchers']:
                body += f"<li style='color: orange;'>⚠️ {voucher['type'].upper()} - {voucher['voucher_number']}: {voucher['status']}</li>"
            
            if not missing_data['incomplete_vouchers']:
                body += "<li style='color: green;'>✅ All vouchers confirmed</li>"
            
            body += """
                </ul>
                
                <h3>Missing Data:</h3>
                <ul>
            """
            
            for data in missing_data['missing_data']:
                body += f"<li style='color: red;'>❌ {data}</li>"
            
            if not missing_data['missing_data']:
                body += "<li style='color: green;'>✅ All data complete</li>"
            
            body += f"""
                </ul>
                
                <div style="margin-top: 30px; padding: 15px; background-color: {'#ffcccc' if urgency == 'URGENT' else '#fff3cd'}; border: 1px solid {'red' if urgency == 'URGENT' else 'orange'};">
                    <p><strong>Action Required:</strong></p>
                    <p>Please update the missing information in the system as soon as possible.</p>
                    <p>Next reminder will be sent: <strong>{frequency.value.replace('_', ' ').title()}</strong></p>
                </div>
                
                <p style="margin-top: 30px; color: #666; font-size: 12px;">
                    This is an automated reminder from Spirit Tours Group Management System.
                </p>
            </body>
            </html>
            """
            
            # Get recipients
            recipients = group.custom_reminder_emails or []
            if not recipients:
                recipients = ['admin@spirittours.com']  # Default admin email
            
            # Send email
            await self.notification_service.send_email(
                to_emails=recipients,
                subject=subject,
                body=body,
                is_html=True,
                priority=urgency
            )
            
            logger.info(f"Sent {urgency} reminder for group {group.group_number}")
            
        except Exception as e:
            logger.error(f"Error sending reminder: {str(e)}")
    
    async def _generate_group_number(self) -> str:
        """Generate unique group number"""
        prefix = "GRP"
        year = datetime.now().year
        
        # Get last group number from database
        last_number = await self._get_last_group_number(year)
        new_number = last_number + 1
        
        return f"{prefix}-{year}-{new_number:05d}"
    
    async def _get_last_group_number(self, year: int) -> int:
        """Get the last group number for the year"""
        # This would query the database for the last group number
        # For now, returning a mock value
        return 42
    
    async def _save_group_coordination(self, group: GroupCoordination):
        """Save group coordination to database"""
        # Save to database implementation
        group.last_updated = datetime.utcnow()
        
        # Invalidate cache
        if self.redis:
            await self.redis.delete(f"group:{group.group_id}")
    
    async def _cache_group(self, group: GroupCoordination):
        """Cache group coordination data"""
        if self.redis:
            cache_key = f"group:{group.group_id}"
            cache_data = pickle.dumps(asdict(group))
            await self.redis.setex(cache_key, self.cache_ttl, cache_data)
    
    async def get_group_coordination(self, group_id: str) -> Optional[GroupCoordination]:
        """Get group coordination by ID"""
        # Check cache first
        if self.redis:
            cache_key = f"group:{group_id}"
            cached = await self.redis.get(cache_key)
            if cached:
                data = pickle.loads(cached)
                return GroupCoordination(**data)
        
        # Get from database if not in cache
        # Database implementation here
        
        return None
    
    async def _update_group_status(self, group: GroupCoordination):
        """Update group status based on assignments and vouchers"""
        has_guide = group.guide_assignment and group.guide_assignment.status == AssignmentStatus.CONFIRMED
        has_driver = group.driver_assignment and group.driver_assignment.status == AssignmentStatus.CONFIRMED
        has_coordinator = group.coordinator_assignment and group.coordinator_assignment.status == AssignmentStatus.CONFIRMED
        
        all_vouchers_confirmed = all(
            v.status == VoucherStatus.CONFIRMED 
            for v in (group.hotel_vouchers + group.restaurant_vouchers + group.entrance_vouchers)
        )
        
        if has_guide and has_driver and has_coordinator and all_vouchers_confirmed:
            group.status = GroupStatus.CONFIRMED
        elif has_guide or has_driver or has_coordinator:
            group.status = GroupStatus.PARTIALLY_ASSIGNED
        else:
            group.status = GroupStatus.PENDING_ASSIGNMENT
    
    async def _check_and_clear_reminders(self, group: GroupCoordination):
        """Check if all requirements are met and clear reminders"""
        missing_data = await self.check_missing_assignments(group.group_id)
        
        if not missing_data['missing_assignments'] and not missing_data['incomplete_vouchers'] and not missing_data['missing_data']:
            # All data complete, send confirmation
            await self._send_completion_notification(group)
            group.reminder_enabled = False
            await self._save_group_coordination(group)
    
    async def _send_assignment_confirmation(self, assignment: GuideAssignment):
        """Send confirmation request to assigned person"""
        # Implementation for sending confirmation request
        pass
    
    async def _send_completion_notification(self, group: GroupCoordination):
        """Send notification that all group data is complete"""
        subject = f"✅ Group {group.group_number} - All Data Complete"
        body = f"""
        <html>
        <body>
            <h2 style="color: green;">Group {group.group_number} is Fully Configured</h2>
            <p>All required assignments and vouchers have been confirmed.</p>
            <p>The group is ready for travel.</p>
        </body>
        </html>
        """
        
        recipients = group.custom_reminder_emails or ['admin@spirittours.com']
        await self.notification_service.send_email(recipients, subject, body, is_html=True)
    
    async def _prepare_report_data(self, group: GroupCoordination, config: ReportConfiguration) -> Dict[str, Any]:
        """Prepare data for report generation based on configuration"""
        report_data = {
            'report_id': config.report_id,
            'report_name': config.report_name,
            'generated_date': datetime.utcnow().isoformat(),
            'group': {}
        }
        
        # Add group info if included
        if config.include_group_info:
            report_data['group']['info'] = {
                'group_number': group.group_number,
                'group_name': group.group_name,
                'status': group.status.value,
                'travel_dates': {
                    'from': group.travel_date_from.isoformat() if group.travel_date_from else None,
                    'to': group.travel_date_to.isoformat() if group.travel_date_to else None
                },
                'total_participants': group.total_participants,
                'group_type': group.group_type
            }
        
        # Add participants if included
        if config.include_participants:
            report_data['group']['participants'] = group.participants
        
        # Add other sections based on configuration
        if config.include_flights:
            report_data['group']['flights'] = [asdict(f) for f in group.flights]
        
        if config.include_hotels:
            report_data['group']['hotels'] = [asdict(h) for h in group.hotel_vouchers]
        
        if config.include_rooming_list:
            report_data['group']['rooming_lists'] = [h.rooming_list for h in group.hotel_vouchers]
        
        # Apply filtering if specified
        if config.date_from or config.date_to:
            # Filter data by date range
            pass
        
        # Apply sorting if specified
        if config.sort_by:
            # Sort data
            pass
        
        return report_data
    
    async def _send_report_email(self, report: bytes, config: ReportConfiguration, group: GroupCoordination):
        """Send generated report via email"""
        subject = f"Report: {config.report_name} - Group {group.group_number}"
        body = f"Please find attached the requested report for group {group.group_number}."
        
        attachment_name = f"{config.report_name}_{group.group_number}.{config.format.value}"
        
        await self.notification_service.send_email_with_attachment(
            to_emails=config.email_to,
            subject=subject,
            body=body,
            attachment=report,
            attachment_name=attachment_name
        )


class VoucherGenerator:
    """Generate vouchers with QR codes and barcodes"""
    
    async def generate_voucher_number(self, voucher_type: VoucherType) -> str:
        """Generate unique voucher number"""
        prefix_map = {
            VoucherType.HOTEL: "HTL",
            VoucherType.RESTAURANT: "RST",
            VoucherType.ENTRANCE_TICKET: "TKT",
            VoucherType.TRANSPORT: "TRN",
            VoucherType.ACTIVITY: "ACT",
            VoucherType.OTHER: "VCH"
        }
        
        prefix = prefix_map.get(voucher_type, "VCH")
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        random_suffix = uuid.uuid4().hex[:6].upper()
        
        return f"{prefix}-{timestamp}-{random_suffix}"
    
    async def generate_qr_code(self, voucher: Voucher) -> str:
        """Generate QR code for voucher"""
        qr_data = {
            'voucher_number': voucher.voucher_number,
            'type': voucher.type.value,
            'group_id': voucher.group_id,
            'amount': voucher.total_amount,
            'currency': voucher.currency,
            'valid_until': voucher.expiry_date.isoformat() if voucher.expiry_date else None
        }
        
        qr = qrcode.QRCode(version=1, box_size=10, border=5)
        qr.add_data(json.dumps(qr_data))
        qr.make(fit=True)
        
        img = qr.make_image(fill_color="black", back_color="white")
        
        # Convert to base64
        buffer = io.BytesIO()
        img.save(buffer, format='PNG')
        img_str = base64.b64encode(buffer.getvalue()).decode()
        
        return f"data:image/png;base64,{img_str}"
    
    async def generate_barcode(self, voucher: Voucher) -> str:
        """Generate barcode for voucher"""
        code = Code128(voucher.voucher_number, writer=ImageWriter())
        
        # Generate barcode image
        buffer = io.BytesIO()
        code.write(buffer)
        
        # Convert to base64
        img_str = base64.b64encode(buffer.getvalue()).decode()
        
        return f"data:image/png;base64,{img_str}"


class ReportGenerator:
    """Generate various report formats"""
    
    async def generate_pdf_report(self, data: Dict[str, Any], config: ReportConfiguration) -> bytes:
        """Generate PDF report"""
        buffer = io.BytesIO()
        
        # Determine page size and orientation
        if config.orientation == "landscape":
            pagesize = landscape(A4 if config.paper_size == "A4" else letter)
        else:
            pagesize = A4 if config.paper_size == "A4" else letter
        
        doc = SimpleDocTemplate(buffer, pagesize=pagesize)
        story = []
        styles = getSampleStyleSheet()
        
        # Add title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1e3a8a'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        story.append(Paragraph(config.report_name, title_style))
        story.append(Spacer(1, 20))
        
        # Add group information
        if config.include_group_info and 'info' in data.get('group', {}):
            info = data['group']['info']
            story.append(Paragraph("Group Information", styles['Heading2']))
            
            info_data = [
                ['Group Number:', info.get('group_number', '')],
                ['Group Name:', info.get('group_name', '')],
                ['Status:', info.get('status', '')],
                ['Travel Dates:', f"{info.get('travel_dates', {}).get('from', '')} to {info.get('travel_dates', {}).get('to', '')}"],
                ['Total Participants:', str(info.get('total_participants', 0))],
                ['Group Type:', info.get('group_type', '')]
            ]
            
            info_table = Table(info_data, colWidths=[150, 350])
            info_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(info_table)
            story.append(Spacer(1, 20))
        
        # Add other sections based on configuration
        # ... (implement other sections)
        
        # Build PDF
        doc.build(story)
        
        return buffer.getvalue()
    
    async def generate_excel_report(self, data: Dict[str, Any], config: ReportConfiguration) -> bytes:
        """Generate Excel report"""
        buffer = io.BytesIO()
        
        # Create workbook
        wb = openpyxl.Workbook()
        
        # Remove default sheet
        wb.remove(wb.active)
        
        # Create sheets based on configuration
        if config.include_group_info:
            ws = wb.create_sheet("Group Information")
            self._add_group_info_to_excel(ws, data.get('group', {}).get('info', {}))
        
        if config.include_participants:
            ws = wb.create_sheet("Participants")
            self._add_participants_to_excel(ws, data.get('group', {}).get('participants', []))
        
        if config.include_flights:
            ws = wb.create_sheet("Flights")
            self._add_flights_to_excel(ws, data.get('group', {}).get('flights', []))
        
        if config.include_hotels:
            ws = wb.create_sheet("Hotels")
            self._add_hotels_to_excel(ws, data.get('group', {}).get('hotels', []))
        
        # Save workbook
        wb.save(buffer)
        
        return buffer.getvalue()
    
    async def generate_word_report(self, data: Dict[str, Any], config: ReportConfiguration) -> bytes:
        """Generate Word document report"""
        doc = Document()
        
        # Add title
        doc.add_heading(config.report_name, 0)
        
        # Add metadata
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("")
        
        # Add group information
        if config.include_group_info and 'info' in data.get('group', {}):
            doc.add_heading("Group Information", 1)
            info = data['group']['info']
            
            table = doc.add_table(rows=6, cols=2)
            table.style = 'Table Grid'
            
            table.cell(0, 0).text = "Group Number:"
            table.cell(0, 1).text = info.get('group_number', '')
            
            table.cell(1, 0).text = "Group Name:"
            table.cell(1, 1).text = info.get('group_name', '')
            
            table.cell(2, 0).text = "Status:"
            table.cell(2, 1).text = info.get('status', '')
            
            table.cell(3, 0).text = "Travel Dates:"
            travel_dates = info.get('travel_dates', {})
            table.cell(3, 1).text = f"{travel_dates.get('from', '')} to {travel_dates.get('to', '')}"
            
            table.cell(4, 0).text = "Total Participants:"
            table.cell(4, 1).text = str(info.get('total_participants', 0))
            
            table.cell(5, 0).text = "Group Type:"
            table.cell(5, 1).text = info.get('group_type', '')
            
            doc.add_paragraph("")
        
        # Add other sections based on configuration
        # ... (implement other sections)
        
        # Save document
        buffer = io.BytesIO()
        doc.save(buffer)
        
        return buffer.getvalue()
    
    async def generate_html_report(self, data: Dict[str, Any], config: ReportConfiguration) -> bytes:
        """Generate HTML report"""
        html_template = """
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{{ report_name }}</title>
            <style>
                body {
                    font-family: Arial, sans-serif;
                    margin: 20px;
                    background-color: #f5f5f5;
                }
                .container {
                    max-width: 1200px;
                    margin: 0 auto;
                    background-color: white;
                    padding: 30px;
                    box-shadow: 0 0 10px rgba(0,0,0,0.1);
                }
                h1 {
                    color: #1e3a8a;
                    border-bottom: 3px solid #1e3a8a;
                    padding-bottom: 10px;
                }
                h2 {
                    color: #334155;
                    margin-top: 30px;
                }
                table {
                    width: 100%;
                    border-collapse: collapse;
                    margin-top: 20px;
                }
                th {
                    background-color: #1e3a8a;
                    color: white;
                    padding: 12px;
                    text-align: left;
                }
                td {
                    padding: 10px;
                    border-bottom: 1px solid #ddd;
                }
                tr:nth-child(even) {
                    background-color: #f9f9f9;
                }
                .info-grid {
                    display: grid;
                    grid-template-columns: 200px 1fr;
                    gap: 10px;
                    margin-top: 20px;
                }
                .info-label {
                    font-weight: bold;
                    color: #555;
                }
                .info-value {
                    color: #333;
                }
                @media print {
                    .container {
                        box-shadow: none;
                    }
                }
            </style>
        </head>
        <body>
            <div class="container">
                <h1>{{ report_name }}</h1>
                <p>Generated: {{ generated_date }}</p>
                
                {% if include_group_info %}
                <h2>Group Information</h2>
                <div class="info-grid">
                    <div class="info-label">Group Number:</div>
                    <div class="info-value">{{ group.info.group_number }}</div>
                    
                    <div class="info-label">Group Name:</div>
                    <div class="info-value">{{ group.info.group_name }}</div>
                    
                    <div class="info-label">Status:</div>
                    <div class="info-value">{{ group.info.status }}</div>
                    
                    <div class="info-label">Travel Dates:</div>
                    <div class="info-value">{{ group.info.travel_dates.from }} to {{ group.info.travel_dates.to }}</div>
                    
                    <div class="info-label">Total Participants:</div>
                    <div class="info-value">{{ group.info.total_participants }}</div>
                    
                    <div class="info-label">Group Type:</div>
                    <div class="info-value">{{ group.info.group_type }}</div>
                </div>
                {% endif %}
                
                <!-- Add other sections here -->
                
            </div>
        </body>
        </html>
        """
        
        template = Template(html_template)
        html_content = template.render(
            report_name=config.report_name,
            generated_date=datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            include_group_info=config.include_group_info,
            group=data.get('group', {})
        )
        
        return html_content.encode('utf-8')
    
    async def generate_json_report(self, data: Dict[str, Any], config: ReportConfiguration) -> bytes:
        """Generate JSON report"""
        # Filter data based on configuration
        filtered_data = {
            'report_metadata': {
                'report_id': config.report_id,
                'report_name': config.report_name,
                'generated_date': datetime.utcnow().isoformat(),
                'format': config.format.value
            }
        }
        
        if config.include_group_info:
            filtered_data['group_info'] = data.get('group', {}).get('info', {})
        
        if config.include_participants:
            filtered_data['participants'] = data.get('group', {}).get('participants', [])
        
        # Add other sections based on configuration
        
        return json.dumps(filtered_data, indent=2, default=str).encode('utf-8')
    
    def _add_group_info_to_excel(self, ws, info: Dict[str, Any]):
        """Add group information to Excel worksheet"""
        ws.append(['Group Information'])
        ws.append([])
        ws.append(['Group Number:', info.get('group_number', '')])
        ws.append(['Group Name:', info.get('group_name', '')])
        ws.append(['Status:', info.get('status', '')])
        ws.append(['Travel Dates:', f"{info.get('travel_dates', {}).get('from', '')} to {info.get('travel_dates', {}).get('to', '')}"])
        ws.append(['Total Participants:', info.get('total_participants', 0)])
        ws.append(['Group Type:', info.get('group_type', '')])
        
        # Apply formatting
        for row in ws.iter_rows(min_row=1, max_row=1):
            for cell in row:
                cell.font = Font(bold=True, size=14)
        
        for row in ws.iter_rows(min_row=3, max_row=8, max_col=1):
            for cell in row:
                cell.font = Font(bold=True)
    
    def _add_participants_to_excel(self, ws, participants: List[Dict[str, Any]]):
        """Add participants to Excel worksheet"""
        if not participants:
            ws.append(['No participants data available'])
            return
        
        # Add headers
        headers = list(participants[0].keys()) if participants else []
        ws.append(headers)
        
        # Add data
        for participant in participants:
            ws.append([participant.get(h, '') for h in headers])
        
        # Apply formatting
        for cell in ws[1]:
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="1e3a8a", end_color="1e3a8a", fill_type="solid")
            cell.font = Font(color="FFFFFF", bold=True)
    
    def _add_flights_to_excel(self, ws, flights: List[Dict[str, Any]]):
        """Add flight information to Excel worksheet"""
        # Implementation for adding flights to Excel
        pass
    
    def _add_hotels_to_excel(self, ws, hotels: List[Dict[str, Any]]):
        """Add hotel information to Excel worksheet"""
        # Implementation for adding hotels to Excel
        pass


class NotificationService:
    """Service for sending notifications and reminders"""
    
    async def send_email(self, to_emails: List[str], subject: str, body: str, 
                        is_html: bool = False, priority: str = "STANDARD") -> bool:
        """Send email notification"""
        try:
            # Email sending implementation
            logger.info(f"Sending email to {to_emails} with subject: {subject}")
            return True
        except Exception as e:
            logger.error(f"Error sending email: {str(e)}")
            return False
    
    async def send_email_with_attachment(self, to_emails: List[str], subject: str, 
                                        body: str, attachment: bytes, 
                                        attachment_name: str) -> bool:
        """Send email with attachment"""
        try:
            # Email with attachment implementation
            logger.info(f"Sending email with attachment to {to_emails}")
            return True
        except Exception as e:
            logger.error(f"Error sending email with attachment: {str(e)}")
            return False


# Export the main system class
__all__ = ['GroupCoordinationSystem', 'GroupCoordination', 'ReportConfiguration', 
           'VoucherType', 'GroupStatus', 'ReportFormat']